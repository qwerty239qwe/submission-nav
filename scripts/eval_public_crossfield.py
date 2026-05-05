from __future__ import annotations

import argparse
from datetime import date
import json
import subprocess
import time
from pathlib import Path

import httpx
from docx import Document

from sn_lib.config import Config
from sn_lib.eval_quality import summarize_rank_quality


OPENALEX_WORKS = "https://api.openalex.org/works"
OPENALEX_SOURCES = "https://api.openalex.org/sources"


FIELDS = {
    "life_sciences_medicine": "clinical disease medicine genomics patient",
    "engineering": "engineering materials structural mechanical manufacturing",
    "computer_science": "computer science machine learning algorithm neural",
    "environmental_sciences": "environmental science climate ecology biodiversity",
    "social_sciences": "social science education inequality psychology policy",
    "chemistry": "chemistry catalysis synthesis molecule organic",
    "physics": "physics quantum condensed matter optical",
}


FIELD_REQUIRED = {
    "life_sciences_medicine": ("clinical", "patient", "medicine", "disease", "cancer", "genom"),
    "engineering": ("engineering", "materials", "structural", "mechanical", "manufacturing", "bridge"),
    "computer_science": ("computer", "algorithm", "machine learning", "neural", "software", "network"),
    "environmental_sciences": ("environmental", "climate", "ecology", "biodiversity", "ecosystem"),
    "social_sciences": ("social", "education", "inequality", "psychology", "policy", "survey"),
    "chemistry": ("chemistry", "chemical", "catalysis", "synthesis", "molecule", "organic"),
    "physics": ("physics", "quantum", "optical", "condensed matter", "particle", "electron"),
}


FIELD_TOPIC_TERMS = {
    "life_sciences_medicine": ("medicine", "clinical medicine", "health sciences", "biology", "genetics"),
    "engineering": ("engineering", "materials science", "mechanical engineering", "electrical engineering"),
    "computer_science": ("computer science", "artificial intelligence", "machine learning", "software"),
    "environmental_sciences": ("environmental science", "earth sciences", "ecology", "climate"),
    "social_sciences": ("social sciences", "psychology", "economics", "education", "political science"),
    "chemistry": ("chemistry", "chemical", "materials chemistry", "organic chemistry"),
    "physics": ("physics", "astronomy", "condensed matter", "optics", "quantum"),
}


TIERS = {
    "top": (5.0, None),
    "middle": (1.5, 5.0),
    "low": (0.0, 1.5),
}


def default_from_date(today: date | None = None) -> str:
    today = today or date.today()
    return date(today.year - 5, 1, 1).isoformat()


def default_to_date(today: date | None = None) -> str:
    return (today or date.today()).isoformat()


def _abstract(index: dict | None) -> str:
    if not index:
        return ""
    words: list[tuple[int, str]] = []
    for word, positions in index.items():
        for pos in positions:
            words.append((pos, word))
    return " ".join(word for _, word in sorted(words))


def _source(work: dict) -> dict:
    return (((work.get("primary_location") or {}).get("source")) or {})


def _impact(work: dict) -> float:
    stats = _source(work).get("summary_stats") or {}
    value = stats.get("2yr_mean_citedness")
    return float(value or 0.0)


def _source_openalex_id(work: dict) -> str | None:
    source_id = _source(work).get("id")
    if not source_id:
        return None
    return source_id.rstrip("/").split("/")[-1]


def _normalize_name(name: str | None) -> str:
    return " ".join((name or "").casefold().replace("&", "and").split())


def _rank_of_name(rows: list[dict], target: str | None) -> int | None:
    target_norm = _normalize_name(target)
    if not target_norm:
        return None
    for idx, row in enumerate(rows, start=1):
        if _normalize_name(row.get("journal")) == target_norm:
            return idx
    return None


def enrich_source_impacts(works: list[dict]) -> None:
    cache: dict[str, dict] = {}
    for work in works:
        source = _source(work)
        stats = source.get("summary_stats") or {}
        if stats.get("2yr_mean_citedness") is not None:
            continue
        source_id = _source_openalex_id(work)
        if not source_id:
            continue
        if source_id not in cache:
            try:
                data = httpx.get(f"{OPENALEX_SOURCES}/{source_id}", timeout=20).json()
            except httpx.HTTPError:
                data = {}
            cache[source_id] = data
        data = cache[source_id]
        if data:
            source["summary_stats"] = data.get("summary_stats") or {}
            source["display_name"] = source.get("display_name") or data.get("display_name")
            source["host_organization_name"] = source.get("host_organization_name") or data.get("host_organization_name")


def _in_tier(impact: float, tier: str) -> bool:
    low, high = TIERS[tier]
    if high is None:
        return impact >= low
    return low <= impact < high


def fetch_candidates(
    query: str,
    per_page: int = 200,
    from_date: str | None = None,
    to_date: str | None = None,
) -> list[dict]:
    filters = ["type:article", "has_abstract:true", "primary_location.source.type:journal"]
    if from_date:
        filters.append(f"from_publication_date:{from_date}")
    if to_date:
        filters.append(f"to_publication_date:{to_date}")
    params = {
        "search": query,
        "filter": ",".join(filters),
        "per-page": per_page,
        "sort": "cited_by_count:desc",
    }
    openalex_mailto = Config.load().key("openalex_email")
    if openalex_mailto:
        params["mailto"] = openalex_mailto
    try:
        response = httpx.get(OPENALEX_WORKS, params=params, timeout=30)
    except httpx.HTTPError as exc:
        raise RuntimeError(f"OpenAlex works fetch failed: {type(exc).__name__}: {exc}") from exc
    if response.status_code != 200:
        message = response.text[:500]
        try:
            payload = response.json()
            message = payload.get("message") or payload.get("error") or message
        except ValueError:
            pass
        raise RuntimeError(f"OpenAlex works fetch failed: HTTP {response.status_code}: {message}")
    data = response.json()
    works = data.get("results") or []
    enrich_source_impacts(works)
    return works


def _matches_field(work: dict, field: str) -> bool:
    haystack = " ".join([
        work.get("title") or "",
        _abstract(work.get("abstract_inverted_index")),
    ]).casefold()
    topic_text = _topic_text(work)
    text_match = any(term in haystack for term in FIELD_REQUIRED[field])
    topic_match = any(term in topic_text for term in FIELD_TOPIC_TERMS[field])
    return topic_match or (text_match and not topic_text)


def _topic_text(work: dict) -> str:
    parts: list[str] = []
    primary = work.get("primary_topic") or {}
    for path in (
        ("display_name",),
        ("subfield", "display_name"),
        ("field", "display_name"),
        ("domain", "display_name"),
    ):
        cur = primary
        for key in path:
            if not isinstance(cur, dict):
                cur = None
                break
            cur = cur.get(key)
        if cur:
            parts.append(str(cur))
    for topic in work.get("topics") or []:
        if topic.get("display_name"):
            parts.append(topic["display_name"])
        for key in ("subfield", "field", "domain"):
            node = topic.get(key) or {}
            if node.get("display_name"):
                parts.append(node["display_name"])
    return " ".join(parts).casefold()


def pick_work(works: list[dict], field: str, tier: str) -> dict | None:
    for work in works:
        source = _source(work)
        if not source.get("display_name"):
            continue
        if not _abstract(work.get("abstract_inverted_index")):
            continue
        if not _matches_field(work, field):
            continue
        if _in_tier(_impact(work), tier):
            return work
    return None


def write_docx(work: dict, path: Path) -> None:
    doc = Document()
    doc.add_heading(work.get("title") or "Untitled", 0)
    authors = []
    for authorship in work.get("authorships") or []:
        author = authorship.get("author") or {}
        if author.get("display_name"):
            authors.append(author["display_name"])
        if len(authors) >= 6:
            break
    if authors:
        doc.add_paragraph(", ".join(authors))
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(_abstract(work.get("abstract_inverted_index")))
    doc.add_heading("Methods", level=1)
    doc.add_paragraph("Public OpenAlex evaluation surrogate generated from title and abstract metadata.")
    doc.add_heading("Results", level=1)
    doc.add_paragraph("This file is used only to test submission-nav venue recommendation behavior.")
    doc.add_heading("References", level=1)
    for idx, ref in enumerate((work.get("referenced_works") or [])[:20], start=1):
        doc.add_paragraph(f"{idx}. {ref}")
    doc.save(str(path))


def run_skill(repo: Path, manuscript: Path, run_dir: Path) -> tuple[int, float, str, str]:
    cmd = [
        "uv", "run", "--project", "scripts", "sn", "strategist",
        str(manuscript),
        "--strategy", "balanced",
        "--venue-types", "journal",
        "--agent-top-n", "15",
        "--run-dir", str(run_dir),
        "--force",
        "--quiet",
    ]
    started = time.perf_counter()
    proc = subprocess.run(cmd, cwd=repo, text=True, capture_output=True, timeout=420)
    elapsed = time.perf_counter() - started
    return proc.returncode, elapsed, proc.stdout, proc.stderr


def summarize_run(run_dir: Path, work: dict, field: str, tier: str, elapsed: float, returncode: int, stderr: str) -> dict:
    ranked_path = run_dir / "ranked_agent_balanced.json"
    raw_ranked_path = run_dir / "ranked_balanced.json"
    contribution_path = run_dir / "contribution_assessment.json"
    ranked = json.loads(ranked_path.read_text(encoding="utf-8")) if ranked_path.exists() else {}
    raw_ranked = json.loads(raw_ranked_path.read_text(encoding="utf-8")) if raw_ranked_path.exists() else []
    contribution = json.loads(contribution_path.read_text(encoding="utf-8")) if contribution_path.exists() else {}
    source = _source(work)
    published = source.get("display_name")
    best_fit = ranked.get("best_fit_ranked") or []
    if not best_fit and raw_ranked:
        from sn_lib.venues import VenueHit
        from sn_lib.ranking import Ranked

        best_fit = [
            Ranked(VenueHit(**row["venue"]), row["score"], row.get("rationale") or {}).to_summary_dict()
            for row in raw_ranked[:15]
        ]
    bucketed_top = ranked.get("risk_adjusted_recommendations") or ranked.get("top") or []
    top_names = [row.get("journal") for row in best_fit[:10]]
    bucketed_names = [row.get("journal") for row in bucketed_top[:10]]
    best_fit_rank = _rank_of_name(best_fit, published)
    bucketed_rank = _rank_of_name(bucketed_top, published)
    full_rank = None
    if raw_ranked:
        full_rows = [{"journal": row.get("venue", {}).get("name")} for row in raw_ranked]
        full_rank = _rank_of_name(full_rows, published)
    demotion_reason = None
    if best_fit_rank is not None and bucketed_rank is None:
        matched = next((row for row in best_fit if _normalize_name(row.get("journal")) == _normalize_name(published)), {})
        demotion_reason = matched.get("bucket_reason") or matched.get("ambition_reason")
    top5_quality = summarize_rank_quality(best_fit, top_n=5)
    top10_quality = summarize_rank_quality(best_fit, top_n=10)
    return {
        "field": field,
        "tier_proxy": tier,
        "paper_title": work.get("title"),
        "publication_year": work.get("publication_year"),
        "publication_date": work.get("publication_date"),
        "published_venue": published,
        "published_venue_impact_proxy": _impact(work),
        "published_venue_in_top10": best_fit_rank is not None and best_fit_rank <= 10,
        "published_best_fit_rank": best_fit_rank,
        "published_full_rank": full_rank,
        "published_bucketed_rank": bucketed_rank,
        "published_bucketed_in_top10": published in bucketed_names,
        "published_demotion_reason": demotion_reason,
        "top_recommendations": top_names,
        "bucketed_recommendations": bucketed_names,
        "top5_quality": top5_quality,
        "top10_quality": top10_quality,
        "contribution_tier": contribution.get("contribution_tier"),
        "ambition_band": contribution.get("ambition_band"),
        "bucket_counts": ranked.get("counts"),
        "elapsed_seconds": round(elapsed, 2),
        "llm_tokens_used_by_cli": 0,
        "artifact_chars": sum(path.stat().st_size for path in run_dir.glob("*.json") if path.is_file()),
        "returncode": returncode,
        "stderr": stderr[-1000:],
    }


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--out-dir", default="temp_sn/public_crossfield_eval")
    ap.add_argument("--fields", nargs="*", default=list(FIELDS))
    ap.add_argument("--tiers", nargs="*", default=list(TIERS))
    ap.add_argument("--from-date", default=default_from_date(),
                    help="Earliest publication date for sampled OpenAlex works. Defaults to January 1 five years ago.")
    ap.add_argument("--to-date", default=default_to_date(),
                    help="Latest publication date for sampled OpenAlex works. Defaults to today.")
    args = ap.parse_args()

    repo = Path(__file__).resolve().parents[1]
    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    manifest: list[dict] = []
    results: list[dict] = []

    def flush_outputs() -> None:
        (out_dir / "results.json").write_text(json.dumps(results, indent=2, ensure_ascii=False), encoding="utf-8")
        (out_dir / "manifest.json").write_text(json.dumps(manifest, indent=2, ensure_ascii=False), encoding="utf-8")

    for field in args.fields:
        try:
            works = fetch_candidates(FIELDS[field], from_date=args.from_date, to_date=args.to_date)
        except RuntimeError as exc:
            for tier in args.tiers:
                results.append({
                    "field": field,
                    "tier_proxy": tier,
                    "from_date": args.from_date,
                    "to_date": args.to_date,
                    "error": "fetch_candidates_failed",
                    "stderr": str(exc)[:1000],
                })
            flush_outputs()
            continue
        for tier in args.tiers:
            work = pick_work(works, field, tier)
            if not work:
                results.append({
                    "field": field,
                    "tier_proxy": tier,
                    "from_date": args.from_date,
                    "to_date": args.to_date,
                    "error": "no OpenAlex work found for tier proxy",
                })
                flush_outputs()
                continue
            case_dir = out_dir / f"{field}_{tier}"
            case_dir.mkdir(parents=True, exist_ok=True)
            manuscript = case_dir / "paper.docx"
            run_dir = case_dir / "run"
            run_dir.mkdir(parents=True, exist_ok=True)
            write_docx(work, manuscript)
            returncode, elapsed, stdout, stderr = run_skill(repo, manuscript, run_dir)
            (case_dir / "stdout.txt").write_text(stdout, encoding="utf-8")
            (case_dir / "stderr.txt").write_text(stderr, encoding="utf-8")
            manifest.append({
                "field": field,
                "tier_proxy": tier,
                "work_id": work.get("id"),
                "publication_year": work.get("publication_year"),
                "publication_date": work.get("publication_date"),
                "from_date": args.from_date,
                "to_date": args.to_date,
                "manuscript": str(manuscript),
                "run_dir": str(run_dir),
            })
            result = summarize_run(run_dir, work, field, tier, elapsed, returncode, stderr)
            result["from_date"] = args.from_date
            result["to_date"] = args.to_date
            results.append(result)
            flush_outputs()
    flush_outputs()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
