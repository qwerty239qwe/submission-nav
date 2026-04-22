import respx, httpx
from docx import Document
from sn_lib.parse import parse_manuscript
from sn_lib.venues import search_venues
from sn_lib.ranking import rank_venues

@respx.mock
def test_end_to_end_strategy(tmp_path, tmp_config_dir):
    p = tmp_path / "m.docx"
    d = Document()
    d.add_heading("Gradient Methods for Widget Optimization", 0)
    d.add_paragraph("Alice")
    d.add_heading("Abstract", 1)
    d.add_paragraph("We optimize widgets via gradient descent.")
    d.save(str(p))
    ms = parse_manuscript(p)
    respx.get("https://api.openalex.org/sources").mock(
        return_value=httpx.Response(200, json={"results": [
            {"id": "A", "display_name": "J Widget", "issn_l": "1",
             "summary_stats": {"2yr_mean_citedness": 4.0, "h_index": 60},
             "is_oa": True, "apc_usd": 1500, "host_organization_name": "X",
             "x_concepts": [{"display_name": "Widget optimization", "score": 0.9}]},
            {"id": "B", "display_name": "J Astronomy", "issn_l": "2",
             "summary_stats": {"2yr_mean_citedness": 5.0, "h_index": 80},
             "is_oa": False, "apc_usd": None, "host_organization_name": "Y",
             "x_concepts": [{"display_name": "Astronomy", "score": 0.9}]},
        ]})
    )
    respx.get("https://api.openalex.org/works").mock(
        return_value=httpx.Response(200, json={"results": [
            {
                "primary_location": {
                    "source": {
                        "id": "A",
                        "display_name": "J Widget",
                        "issn_l": "1",
                        "is_oa": True,
                        "host_organization_name": "X",
                    }
                },
                "primary_topic": {"display_name": "Widget optimization"},
            }
        ]})
    )
    venues = search_venues(ms.abstract or ms.title, per_page=10)
    ranked = rank_venues(["widget optimization", "gradient methods"], venues)
    assert ranked[0].venue.name == "J Widget"
