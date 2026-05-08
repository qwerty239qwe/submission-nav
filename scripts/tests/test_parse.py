from pathlib import Path
from docx import Document
from sn_lib.parse import (
    parse_manuscript,
    Manuscript,
    _extract_pdf_abstract_from_layout,
    _infer_article_type,
    _pick_pdf_title_and_authors,
    _pick_pdf_title_from_layout,
)

def _make_docx(path: Path):
    d = Document()
    d.add_heading("A Novel Method for Widget Optimization", 0)
    d.add_paragraph("Alice Smith, Bob Jones")
    d.add_heading("Abstract", 1)
    d.add_paragraph("We present a new approach to widget optimization using gradient methods.")
    d.add_heading("Introduction", 1)
    d.add_paragraph("Widgets are important. Prior work ignored X.")
    d.add_heading("References", 1)
    d.add_paragraph("[1] Foo 2020. [2] Bar 2021.")
    d.save(str(path))

def test_parse_docx(tmp_path):
    p = tmp_path / "s.docx"
    _make_docx(p)
    m = parse_manuscript(p)
    assert isinstance(m, Manuscript)
    assert "Widget" in m.title
    assert m.abstract and "widget" in m.abstract.lower()
    assert any("Introduction" in s.heading for s in m.sections)
    assert m.word_count > 10
    assert m.reference_count >= 2

def test_parse_docx_prefers_plain_title_over_section_heading(tmp_path):
    p = tmp_path / "plain-title.docx"
    d = Document()
    d.add_paragraph("A Plain Title Without Title Style")
    d.add_paragraph("Alice Smith, Bob Jones")
    d.add_heading("Abstract", 1)
    d.add_paragraph("This abstract should not become the title.")
    d.save(str(p))
    m = parse_manuscript(p)
    assert m.title == "A Plain Title Without Title Style"
    assert m.authors == ["Alice Smith", "Bob Jones"]


def test_parse_docx_handles_numbered_headings_and_numbered_references(tmp_path):
    p = tmp_path / "numbered.docx"
    d = Document()
    d.add_paragraph("Interpretable Models for Mitochondrial Toxicity")
    d.add_paragraph("Alice Smith, Bob Jones")
    d.add_paragraph("1. Abstract")
    d.add_paragraph("This abstract should stay short.")
    d.add_paragraph("2. Materials and Methods")
    d.add_paragraph("Methods content.")
    d.add_paragraph("3. Results")
    d.add_paragraph("Results content.")
    d.add_paragraph("References")
    d.add_paragraph("1. Foo et al. 2020. https://doi.org/10.1000/foo")
    d.add_paragraph("2. Bar et al. 2021. https://doi.org/10.1000/bar")
    d.save(str(p))
    m = parse_manuscript(p)
    assert m.abstract == "This abstract should stay short."
    assert any("Materials and Methods" in s.heading for s in m.sections)
    assert any("Results" in s.heading for s in m.sections)
    assert m.reference_count == 2
    assert len(m.references) == 2
    assert "10.1000/foo" in m.references[0]


def test_parse_summary_dict_is_compact(tmp_path):
    p = tmp_path / "summary.docx"
    _make_docx(p)
    m = parse_manuscript(p)
    summary = m.to_summary_dict()
    assert summary["title"] == m.title
    assert "section_headings" in summary
    assert "sections" not in summary
    assert summary["reference_count"] == m.reference_count


def test_pick_pdf_title_and_authors_handles_plos_style_front_matter():
    lines = [
        "RESEARCH ARTICLE",
        "Galaxy-ML: An accessible, reproducible, and",
        "scalable machine learning toolkit for biomedicine",
        "Qiang Gu1,2, Anup Kumar3, Simon Bray3",
        "Department of Biomedical Engineering, Oregon Health & Science University",
        "Abstract",
    ]
    title, authors = _pick_pdf_title_and_authors(lines)
    assert title == "Galaxy-ML: An accessible, reproducible, and scalable machine learning toolkit for biomedicine"
    assert authors == ["Qiang Gu1,2, Anup Kumar3, Simon Bray3"]


def test_pick_pdf_title_from_layout_skips_journal_and_article_labels():
    rows = [
        (31.4, 25.0, "Scientific Reports"),
        (48.9, 9.0, "https://doi.org/10.1038/example"),
        (66.5, 21.0, "Article in Press"),
        (114.5, 28.0, "Protium heptaphyllum, a tree native"),
        (145.3, 28.0, "to the Atlantic Forest, is a potential source"),
        (175.3, 28.0, "of compounds against important cocoa"),
        (205.3, 28.0, "phytopathogen"),
    ]
    assert _pick_pdf_title_from_layout(rows) == (
        "Protium heptaphyllum, a tree native to the Atlantic Forest, is a potential source "
        "of compounds against important cocoa phytopathogen"
    )


def test_pick_pdf_title_from_layout_skips_target_journal_header():
    rows = [
        (52.3, 20.9, "ChemComm"),
        (94.2, 9.0, "View Article Online"),
        (95.2, 15.9, "REVIEW ARTICLE"),
        (106.4, 6.0, "View Journal"),
        (136.0, 15.9, "Revealing the reaction pathways and interfacial"),
        (153.9, 15.9, "regulation mechanisms of urea electro-oxidation"),
        (171.8, 15.9, "on nickel-based catalysts"),
        (172.8, 8.0, "Open Access Article. Published on 29 April 2026. Downloaded on 5/7/2026 3:16:12 AM."),
    ]
    assert _pick_pdf_title_from_layout(rows) == (
        "Revealing the reaction pathways and interfacial regulation mechanisms of urea electro-oxidation "
        "on nickel-based catalysts"
    )


def test_pick_pdf_title_from_layout_includes_arxiv_subtitle_not_sidebar():
    rows = [
        (116.1, 20.7, "The Adversarial Discount"),
        (139.0, 14.3, "AI, Signal Correlation, and the Cybersecurity Arms Race"),
        (174.5, 14.3, "James Bono"),
        (202.4, 14.3, "May 7, 2026"),
        (213.2, 20.0, "arXiv:2605.04336v1  [econ.TH]  5 May 2026"),
        (248.6, 10.9, "Abstract"),
    ]
    assert _pick_pdf_title_from_layout(rows) == (
        "The Adversarial Discount AI, Signal Correlation, and the Cybersecurity Arms Race"
    )


def test_extract_pdf_abstract_from_layout_skips_target_journal_metadata():
    title = (
        "Revealing the reaction pathways and interfacial regulation mechanisms of urea electro-oxidation "
        "on nickel-based catalysts"
    )
    rows = [
        (52.3, 20.9, "ChemComm"),
        (136.0, 15.9, "Revealing the reaction pathways and interfacial"),
        (153.9, 15.9, "regulation mechanisms of urea electro-oxidation"),
        (171.8, 15.9, "on nickel-based catalysts"),
        (195.9, 10.0, "Riyi Zhang, ab Yong Yan,* ab Zhihao Yi"),
        (225.8, 8.0, "With an equilibrium potential of 0.37 V versus RHE, much lower than that of the oxygen evolution"),
        (237.8, 8.0, "reaction, the electrocatalytic urea oxidation reaction has been widely regarded as a promising anodic"),
        (249.7, 8.0, "half reaction for reducing the energy consumption of hydrogen production while treating nitrogen"),
        (261.7, 8.0, "containing wastewater. Its practical application, however, is still limited by sluggish kinetics arising from"),
        (273.6, 8.0, "the complex six electron, six proton transfer process, even in high performance nickel based catalysts."),
        (430.1, 7.2, "Received 25th March 2026,"),
    ]
    abstract = _extract_pdf_abstract_from_layout(rows, title)
    assert abstract is not None
    assert abstract.startswith("With an equilibrium potential")
    assert "ChemComm" not in abstract
    assert "Received 25th March" not in abstract


def test_infer_article_type_from_review_article_front_matter():
    assert _infer_article_type(["ChemComm", "REVIEW ARTICLE", "Title"]) == "review"


def test_pick_pdf_title_from_layout_handles_interleaved_sidebar_metadata():
    rows = [
        (50.4, 7.0, "TYPE Original Research"),
        (136.9, 21.0, "The protective effect of"),
        (176.8, 7.0, "OPEN ACCESS"),
        (186.9, 21.0, "probiotics on intestinal"),
        (192.4, 5.5, "EDITED BY"),
        (211.9, 21.0, "mucosal injury and dysbiosis"),
        (236.9, 21.0, "in infants with congenital"),
        (261.9, 21.0, "heart disease undergoing"),
        (300.4, 11.0, "Zhixuan Zhang 1"),
    ]
    assert _pick_pdf_title_from_layout(rows) == (
        "The protective effect of probiotics on intestinal mucosal injury and dysbiosis "
        "in infants with congenital heart disease undergoing"
    )
