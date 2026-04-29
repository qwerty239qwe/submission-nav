import json

import respx
import httpx
from docx import Document

from sn_lib.cli import _concept_queries, main


def _docx(path):
    doc = Document()
    doc.add_heading("Gradient Methods for Widget Optimization", 0)
    doc.add_paragraph("Alice")
    doc.add_heading("Abstract", 1)
    doc.add_paragraph("We optimize widgets via gradient descent.")
    doc.save(str(path))


def test_cli_parse_and_concepts(tmp_path, tmp_config_dir, capsys):
    manuscript = tmp_path / "paper.docx"
    _docx(manuscript)
    assert main(["parse", str(manuscript)]) == 0
    out = capsys.readouterr().out
    assert "OK parse" in out
    assert main(["concepts", str(manuscript)]) == 0
    out = capsys.readouterr().out
    assert "OK concepts" in out


def test_cli_profile_writes_manuscript_profile(tmp_path, tmp_config_dir, capsys):
    manuscript = tmp_path / "paper.docx"
    _docx(manuscript)
    assert main(["profile", str(manuscript)]) == 0
    assert "OK profile" in capsys.readouterr().out
    profile_path = next((tmp_config_dir / "runs").glob("*/ms_profile.json"))
    payload = json.loads(profile_path.read_text(encoding="utf-8"))
    assert payload["contribution_type"] in {"original_research", "clinical"}
    assert "data_type" in payload


def test_cli_contribution_writes_assessment(tmp_path, tmp_config_dir, capsys):
    manuscript = tmp_path / "paper.docx"
    _docx(manuscript)
    assert main(["contribution", str(manuscript)]) == 0
    assert "OK contribution" in capsys.readouterr().out
    contribution_path = next((tmp_config_dir / "runs").glob("*/contribution_assessment.json"))
    payload = json.loads(contribution_path.read_text(encoding="utf-8"))
    assert payload["contribution_tier"] in {"exploratory", "solid_specialty", "strong_specialty", "high_impact_specialty", "elite_general"}
    assert "scores" in payload


def test_cli_config_set_supports_emails(tmp_path, tmp_config_dir, monkeypatch, capsys):
    monkeypatch.setattr("sn_lib.config._dotenv_path", lambda: tmp_path / ".env")
    assert main(["config", "set", "--key", "openalex_email", "--value", "me@example.org"]) == 0
    assert "OPENALEX_EMAIL=me@example.org" in (tmp_path / ".env").read_text(encoding="utf-8")
    assert "OK config" in capsys.readouterr().out
    assert main(["config", "show"]) == 0
    payload = json.loads(capsys.readouterr().out)
    assert payload["openalex_email"] == "me@example.org"


def test_cli_config_set_can_store_config_json(tmp_config_dir):
    assert main(["config", "set", "--key", "crossref_email", "--value", "me@example.org", "--store", "config"]) == 0
    payload = json.loads((tmp_config_dir / "config.json").read_text(encoding="utf-8"))
    assert payload["crossref_email"] == "me@example.org"


def test_cli_rules_from_file(tmp_path, tmp_config_dir, capsys):
    html = tmp_path / "rules.html"
    html.write_text("<p>The abstract should not exceed 250 words.</p>", encoding="utf-8")
    assert main(["rules", "Example Journal", "--from-file", str(html)]) == 0
    out = capsys.readouterr().out
    assert "OK rules" in out
    payload = json.loads((tmp_config_dir / "rules" / "example-journal.json").read_text(encoding="utf-8"))
    assert payload["abstract_limit"] == 250


def test_cli_doctor_reports_publisher_risk_path(tmp_config_dir, capsys):
    assert main(["doctor"]) == 0
    out = capsys.readouterr().out
    payload_text = out.split("\nOK doctor", 1)[0]
    payload = json.loads(payload_text)
    assert payload["publisher_risk_path"].endswith("publisher_risk.json")
    assert payload["publisher_risk_configured"] is False


def test_concept_queries_prefers_stored_queries_and_falls_back_for_old_cache():
    assert _concept_queries({"concepts": ["new widget"], "queries": ["stored query"]}, 4) == ["stored query"]
    assert _concept_queries({"concepts": ["new widget", "gradient descent"]}, 4)[0] == "new widget gradient descent"


@respx.mock
def test_cli_strategist_runs_chained_workflow(tmp_path, tmp_config_dir, monkeypatch, capsys):
    monkeypatch.delenv("ELSEVIER_API_KEY", raising=False)
    monkeypatch.delenv("SCOPUS_API_KEY", raising=False)
    monkeypatch.setattr("sn_lib.config._dotenv_path", lambda: tmp_path / ".env")
    manuscript = tmp_path / "paper.docx"
    _docx(manuscript)
    respx.get("https://api.openalex.org/sources").mock(
        return_value=httpx.Response(200, json={"results": [
            {"id": "A", "display_name": "J Widget", "issn_l": "1",
             "summary_stats": {"2yr_mean_citedness": 4.0, "h_index": 60},
             "is_oa": True, "apc_usd": 1500, "host_organization_name": "X",
             "x_concepts": [{"display_name": "Widget optimization", "score": 0.9}]},
        ]})
    )
    respx.get("https://api.openalex.org/works").mock(
        return_value=httpx.Response(200, json={"results": [
            {
                "primary_location": {
                    "source": {
                        "id": "A",
                        "display_name": "J Widget",
                        "issn_l": "1",
                        "is_oa": True,
                        "host_organization_name": "X",
                    }
                },
                "primary_topic": {"display_name": "Widget optimization"},
            }
        ]})
    )
    respx.get("https://api.crossref.org/journals/1").mock(
        return_value=httpx.Response(200, json={"message": {"title": "J Widget", "publisher": "X", "ISSN": ["1"]}})
    )
    assert main(["strategist", str(manuscript), "--per-page", "5", "--agent-top-n", "3"]) == 0
    out = capsys.readouterr().out
    assert "OK strategist" in out
    assert "J Widget" in out
    payload = json.loads(out[out.index("{"):])
    assert "buckets" in payload
    assert payload["top"][0]["journal"] == "J Widget"
